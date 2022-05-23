import docx
doc = docx.Document()
doc.add_heading('Python', 0)
doc_para = doc.add_paragraph('Introduction to Python , ')
records = (
    ('Introduction to Python'),
    ('Control Statements'),
    ('List, Ranges & Tuples in Python'),
    ("Python Dictionaries and Sets"),
    ("Input and Output in Python"),
    ("Python built in function"),
    ("Python OOPS"),
    ("Exceptions"),
    ("Python Regular Expressions"),
    ("Using Databases in Python")
)

table = doc.add_table(rows=1, cols=1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Topics'

for topics in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(topics)
    #row_cells[1].text = id
    #row_cells[2].text = desc

doc_para.add_run().bold = True
doc_para.add_run()
doc_para.add_run().italic = True

doc.add_page_break()
doc.add_heading('Images', 2)
doc.add_picture("E:\own\Screenshot 2022-03-19 225752.jpg")
doc.add_paragraph("A mental image exists in an individual's mind, as something one remembers or imagines. The subject of an image need not be real; it may be an abstract concept, such as a graph, function, or imaginary entity. For example, Sigmund Freud claimed to have dreamed purely in aural-images of dialogs.[citation needed] Different scholars of psychoanalysis as well as the social sciences such as Slavoj Žižek and Jan Berger have pointed out the possibility of manipulating mental images for ideological purposes. Images perpetuated in public education, media as well as popular culture have a profound impact on the formation of such mental images:")



doc.save('E:/teja.docx')
